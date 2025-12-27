from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---------------- HEADER SECTION ----------------
def add_centered_bold(text, size=12, after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)

add_centered_bold("FORM ‘A’", 12)
add_centered_bold("MEDIATION APPLICATION FORM", 12)
add_centered_bold("[REFER RULE 3(1)]", 12, after=10)

p = doc.add_paragraph("Mumbai District Legal Services Authority")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("City Civil Court, Mumbai")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("\n")

# ---------------- TABLE SETUP ----------------
table = doc.add_table(rows=0, cols=3)
table.style = "Table Grid"
table.autofit = False
table.columns[0].width = Inches(0.4)
table.columns[1].width = Inches(2.5)
table.columns[2].width = Inches(3.8)

def add_row(c1, c2, c3, bold_label=True):
    row = table.add_row()
    row.cells[0].text = c1
    run = row.cells[1].paragraphs[0].add_run(c2)
    run.bold = bold_label
    row.cells[2].text = c3
    return row

# ---------------- DETAILS OF PARTIES ----------------
row = table.add_row()
row.cells[0].merge(row.cells[1]).merge(row.cells[2])
row.cells[0].paragraphs[0].add_run("DETAILS OF PARTIES:").bold = True

add_row("1", "Name of Applicant", "{{client_name}}")

row = table.add_row()
row.cells[0].merge(row.cells[1]).merge(row.cells[2])
row.cells[0].paragraphs[0].add_run(
    "Address and contact details of Applicant"
).bold = True

# -------- Applicant Address (FIXED: multiple paragraphs) --------
row = table.add_row()
row.cells[0].text = "1"
row.cells[1].paragraphs[0].add_run("Address").bold = True

cell = row.cells[2]
cell.text = ""

p1 = cell.add_paragraph()
p1.add_run("REGISTERED ADDRESS:").bold = True

p2 = cell.add_paragraph()
p2.add_run("{{branch_address}}")

p3 = cell.add_paragraph()
p3.add_run("CORRESPONDENCE BRANCH ADDRESS:").bold = True

p4 = cell.add_paragraph()
p4.add_run("{{branch_address}}")

add_row("", "Telephone No.", "{{mobile}}")
add_row("", "Mobile No.", "")
add_row("", "Email ID", "info@kslegal.co.in")

# ---------------- OPPOSITE PARTY ----------------
row = table.add_row()
row.cells[0].merge(row.cells[1]).merge(row.cells[2])
row.cells[0].paragraphs[0].add_run(
    "2  Name, Address and Contact details of Opposite Party:"
).bold = True

row = table.add_row()
row.cells[0].merge(row.cells[1]).merge(row.cells[2])
row.cells[0].paragraphs[0].add_run(
    "Address and contact details of Defendant/s"
).bold = True

add_row("", "Name", "{{customer_name}}")

# -------- Defendant Address (FIXED) --------
row = table.add_row()
row.cells[0].text = ""
row.cells[1].paragraphs[0].add_run("Address").bold = True

cell = row.cells[2]
cell.text = ""

p1 = cell.add_paragraph()
p1.add_run("REGISTERED ADDRESS:").bold = True

p2 = cell.add_paragraph()
p2.add_run(
    "{% if address1 and address1 != \"\" %} {{address1}} {% else %} __________________ {% endif %}"
)

p3 = cell.add_paragraph()
p3.add_run("CORRESPONDENCE ADDRESS:").bold = True

p4 = cell.add_paragraph()
p4.add_run(
    "{% if address1 and address1 != \"\" %} {{address1}} {% else %} __________________ {% endif %}"
)

add_row("", "Telephone No.", "")
add_row("", "Mobile No.", "")
add_row("", "Email ID", "")

# ---------------- DETAILS OF DISPUTE ----------------
row = table.add_row()
row.cells[0].merge(row.cells[1]).merge(row.cells[2])
row.cells[0].paragraphs[0].add_run("DETAILS OF DISPUTE:").bold = True

row = table.add_row()
row.cells[0].merge(row.cells[1]).merge(row.cells[2])
p = row.cells[0].paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "THE COMM. COURTS (PRE-INSTITUTION.........SETTLEMENT) RULES, 2018"
)
run.bold = True
run.underline = True

row = table.add_row()
row.cells[0].merge(row.cells[1]).merge(row.cells[2])
row.cells[0].paragraphs[0].add_run(
    "Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, 2015 (4 of 2016):"
).bold = True


doc.save("output/form_A_corrected.docx")
